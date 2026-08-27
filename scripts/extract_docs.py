# -*- coding: utf-8 -*-
"""提取国别指南 docx 的正文与标题结构，输出结构化 JSON 索引。"""
import os, json, re
from docx import Document

DOCS = {
    "阿联酋": r"C:/Users/admin/.wpscomate/agent/workspace/zhidu-qa/docs/中国海诚境外财务管理操作指南之阿联酋.docx",
    "摩洛哥": r"C:/Users/admin/.wpscomate/agent/workspace/zhidu-qa/docs/中国海诚境外财务管理操作指南之摩洛哥.docx",
}
OUT = r"C:/Users/admin/.wpscomate/agent/workspace/zhidu-qa/data"

def heading_level(style_name):
    s = (style_name or "").lower()
    m = re.search(r"heading\s*(\d+)", s)
    if m:
        return int(m.group(1))
    if "标题" in s:
        m = re.search(r"标题\s*(\d+)", s)
        if m:
            return int(m.group(1))
    return 0

def extract(path):
    doc = Document(path)
    blocks = []  # (level, text)
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        lvl = heading_level(p.style.name if p.style else "")
        blocks.append((lvl, text))
    # 表格内容也提取（指南里可能有表格）
    tables_text = []
    for ti, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            rows.append(" | ".join(cells))
        tables_text.append("\n".join(rows))
    return blocks, tables_text

def build_index(country, blocks, tables_text):
    """把段落按标题组织成章节树，并切分成检索片段。"""
    # 先构建章节树
    root = {"title": country, "level": 0, "children": [], "content": []}
    stack = [root]
    for lvl, text in blocks:
        if lvl >= 1:
            node = {"title": text, "level": lvl, "children": [], "content": []}
            while len(stack) > 1 and stack[-1]["level"] >= lvl:
                stack.pop()
            stack[-1]["children"].append(node)
            stack.append(node)
        else:
            stack[-1]["content"].append(text)
    # 表格作为附录内容
    if tables_text:
        root["content"].append("\n\n【表格内容】\n" + "\n\n".join(tables_text))
    return root

def split_body(body):
    """把节点正文按【表格内容】标记拆成多个块，超长块再按空行切分。"""
    segs = re.split(r'(【表格内容】)', body)
    blocks = []
    cur = ""
    for s in segs:
        if s == '【表格内容】':
            if cur.strip():
                blocks.append(cur.strip())
            cur = ""
        else:
            cur += s
    if cur.strip():
        blocks.append(cur.strip())
    # 超长块按空行二次切分
    result = []
    for blk in blocks:
        if len(blk) <= 1500:
            result.append(blk)
            continue
        subs = re.split(r'\n\s*\n', blk)
        acc = ""
        for s in subs:
            if acc and len(acc) + len(s) > 1200:
                result.append(acc)
                acc = s
            else:
                acc = (acc + "\n" + s) if acc else s
        if acc:
            result.append(acc)
    return result


def flatten(node, country, path, out):
    """深度优先展开，生成检索片段（按章节聚合正文，表格/超长块拆分）。"""
    cur_path = path + [node["title"]]
    body = "\n".join(node["content"]).strip()
    if body:
        blocks = split_body(body)
        for bi, blk in enumerate(blocks):
            suffix = f"（附表{bi+1}）" if len(blocks) > 1 else ""
            out.append({
                "country": country,
                "section": " / ".join(cur_path) + suffix,
                "text": blk,
            })
    for c in node["children"]:
        flatten(c, country, cur_path, out)

def main():
    os.makedirs(OUT, exist_ok=True)
    all_chunks = []
    for country, path in DOCS.items():
        blocks, tables_text = extract(path)
        tree = build_index(country, blocks, tables_text)
        chunks = []
        flatten(tree, country, [], chunks)
        all_chunks.extend(chunks)
        # 保存章节树，便于查看结构
        with open(os.path.join(OUT, f"tree_{country}.json"), "w", encoding="utf-8") as f:
            json.dump(tree, f, ensure_ascii=False, indent=2)
        print(f"[{country}] 段落数={len(blocks)}, 表格数={len(tables_text)}, 检索片段数={len(chunks)}")
        # 打印前 40 个标题
        titles = [b[1] for b in blocks if b[0] >= 1]
        print(f"  标题数={len(titles)}")
        for t in titles[:40]:
            print("   -", t)
    with open(os.path.join(OUT, "chunks.json"), "w", encoding="utf-8") as f:
        json.dump(all_chunks, f, ensure_ascii=False, indent=2)
    print(f"\n总片段数={len(all_chunks)}")
    print(f"OUTPUT={os.path.abspath(os.path.join(OUT, 'chunks.json'))}")

if __name__ == "__main__":
    main()
